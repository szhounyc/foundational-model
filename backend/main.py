from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import httpx
import os
import json
import uuid
from datetime import datetime
from typing import List, Optional, Dict
from pydantic import BaseModel
import sqlite3
from contextlib import asynccontextmanager
import aiofiles
import PyPDF2
import io
import logging
import difflib
import re
from docx import Document
from pathlib import Path
import shutil
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Database configuration
DATABASE_PATH = os.environ.get('DATABASE_PATH', 'app.db')
logger.info(f"Using database path: {DATABASE_PATH}")

# Database models
class TrainingJob(BaseModel):
    id: str
    model_name: str
    dataset_files: List[str]
    status: str
    created_at: str
    completed_at: Optional[str] = None
    hyperparameters: dict
    progress: float = 0.0

class InferenceRequest(BaseModel):
    model_id: str
    prompt: str
    max_tokens: int = 100
    temperature: float = 0.7

class ContractReviewRequest(BaseModel):
    contract_files: List[str]  # List of file IDs
    model_id: str = "sftj-s1xkr35z"
    review_type: str = "comprehensive"

class ModelInfo(BaseModel):
    id: str
    name: str
    type: str
    status: str
    size_mb: float
    created_at: str

# New models for contract comment feature
class LawFirm(BaseModel):
    id: str
    name: str
    keywords: List[str]
    created_at: str

class ContractTemplate(BaseModel):
    id: str
    law_firm_id: str
    law_firm_name: str
    template_type: str  # "main_contract", "rider", "legal_comments"
    file_name: str
    file_path: str
    content: str
    created_at: str

class ContractCommentRequest(BaseModel):
    contract_file_id: str
    rider_file_id: Optional[str] = None
    law_firm_id: Optional[str] = None

class BatchContractCommentRequest(BaseModel):
    contract_file_ids: List[str]
    law_firm_id: Optional[str] = None

class DocumentComparison(BaseModel):
    id: str
    contract_file_id: str
    template_id: str
    differences: List[dict]
    created_at: str

class LawFirmCreate(BaseModel):
    name: str
    keywords: List[str]

# Database initialization
def init_db():
    # Ensure database directory exists
    db_path = Path(DATABASE_PATH)
    db_path.parent.mkdir(parents=True, exist_ok=True)
    
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    # Training jobs table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS training_jobs (
            id TEXT PRIMARY KEY,
            model_name TEXT NOT NULL,
            dataset_files TEXT NOT NULL,
            status TEXT NOT NULL,
            created_at TEXT NOT NULL,
            completed_at TEXT,
            hyperparameters TEXT NOT NULL,
            progress REAL DEFAULT 0.0
        )
    ''')
    
    # Models table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS models (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            type TEXT NOT NULL,
            status TEXT NOT NULL,
            size_mb REAL NOT NULL,
            created_at TEXT NOT NULL,
            file_path TEXT
        )
    ''')
    
    # PDF files table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS pdf_files (
            id TEXT PRIMARY KEY,
            filename TEXT NOT NULL,
            file_path TEXT NOT NULL,
            processed BOOLEAN DEFAULT FALSE,
            uploaded_at TEXT NOT NULL,
            size_bytes INTEGER,
            extracted_text TEXT,
            document_type TEXT DEFAULT NULL
        )
    ''')
    
    # Contract reviews table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS contract_reviews (
            id TEXT PRIMARY KEY,
            contract_files TEXT NOT NULL,
            model_id TEXT NOT NULL,
            review_type TEXT NOT NULL,
            status TEXT NOT NULL,
            created_at TEXT NOT NULL,
            completed_at TEXT,
            review_result TEXT,
            processing_time REAL
        )
    ''')
    
    # Contract comments table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS contract_comments (
            id TEXT PRIMARY KEY,
            contract_file_id TEXT NOT NULL,
            law_firm_id TEXT NOT NULL,
            status TEXT NOT NULL,
            created_at TEXT NOT NULL,
            completed_at TEXT,
            comments_result TEXT,
            processing_time REAL,
            FOREIGN KEY (contract_file_id) REFERENCES pdf_files (id),
            FOREIGN KEY (law_firm_id) REFERENCES law_firms (id)
        )
    ''')
    
    # Law firms table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS law_firms (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL UNIQUE,
            keywords TEXT NOT NULL,
            created_at TEXT NOT NULL
        )
    ''')
    
    # Contract templates table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS contract_templates (
            id TEXT PRIMARY KEY,
            law_firm_id TEXT NOT NULL,
            law_firm_name TEXT NOT NULL,
            template_type TEXT NOT NULL,
            file_name TEXT NOT NULL,
            file_path TEXT NOT NULL,
            content TEXT,
            created_at TEXT NOT NULL,
            FOREIGN KEY (law_firm_id) REFERENCES law_firms (id)
        )
    ''')
    
    # Document comparisons table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS document_comparisons (
            id TEXT PRIMARY KEY,
            contract_file_id TEXT NOT NULL,
            template_id TEXT NOT NULL,
            differences TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY (contract_file_id) REFERENCES pdf_files (id),
            FOREIGN KEY (template_id) REFERENCES contract_templates (id)
        )
    ''')
    
    # Add law_firm_id column to pdf_files table if it doesn't exist
    cursor.execute("PRAGMA table_info(pdf_files)")
    columns = [column[1] for column in cursor.fetchall()]
    if 'law_firm_id' not in columns:
        cursor.execute('ALTER TABLE pdf_files ADD COLUMN law_firm_id TEXT')
    if 'law_firm_name' not in columns:
        cursor.execute('ALTER TABLE pdf_files ADD COLUMN law_firm_name TEXT')
    
    # Insert default law firm
    cursor.execute('''
        INSERT OR IGNORE INTO law_firms (id, name, keywords, created_at)
        VALUES (?, ?, ?, ?)
    ''', (
        'law_firm_1',
        'Marans Newman Tsolis & Nazinitsky LLC',
        json.dumps(['Marans Newman Tsolis', 'MNTN', 'Marans Newman', 'Tsolis', 'Nazinitsky']),
        datetime.now().isoformat()
    ))
    
    conn.commit()
    conn.close()
    logger.info("✅ Database initialized")

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    logger.info(f"Starting PDF text extraction for: {file_path}")
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            page_count = len(pdf_reader.pages)
            logger.info(f"PDF has {page_count} pages")
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text += page_text + "\n"
                logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
            
        total_chars = len(text.strip())
        logger.info(f"Successfully extracted {total_chars} characters from PDF: {file_path}")
        
        # Log first 500 characters for debugging
        preview_text = text.strip()[:500]
        logger.info(f"PDF content preview (first 500 chars): {preview_text}...")
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error extracting text from PDF: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    logger.info(f"Starting DOCX text extraction for: {file_path}")
    try:
        doc = Document(file_path)
        text = ""
        paragraph_count = len(doc.paragraphs)
        logger.info(f"DOCX has {paragraph_count} paragraphs")
        
        for paragraph_num, paragraph in enumerate(doc.paragraphs):
            paragraph_text = paragraph.text
            text += paragraph_text + "\n"
            logger.debug(f"Extracted {len(paragraph_text)} characters from paragraph {paragraph_num + 1}")
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        total_chars = len(text.strip())
        logger.info(f"Successfully extracted {total_chars} characters from DOCX: {file_path}")
        
        # Log first 500 characters for debugging
        preview_text = text.strip()[:500]
        logger.info(f"DOCX content preview (first 500 chars): {preview_text}...")
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error extracting text from DOCX: {str(e)}")

def extract_text_from_file(file_path: str) -> str:
    """Extract text from PDF or DOCX file based on extension"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")

def is_supported_file(filename: str) -> bool:
    """Check if file type is supported (PDF or DOCX)"""
    return filename.lower().endswith(('.pdf', '.docx'))

def detect_law_firm(text: str) -> tuple:
    """Detect law firm from contract text"""
    logger.info("Starting law firm detection...")
    
    # Get all law firms from database
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT id, name, keywords FROM law_firms')
    law_firms = cursor.fetchall()
    conn.close()
    
    text_lower = text.lower()
    
    for firm_id, firm_name, keywords_json in law_firms:
        keywords = json.loads(keywords_json)
        for keyword in keywords:
            if keyword.lower() in text_lower:
                logger.info(f"Law firm detected: {firm_name} (keyword: {keyword})")
                return firm_id, firm_name
    
    logger.info("No law firm detected in contract text")
    return None, None

import difflib
import re

def compare_documents(doc1_text: str, doc2_text: str) -> List[dict]:
    """Compare two documents and return differences"""
    logger.info("Starting document comparison...")
    
    # Split documents into lines for comparison
    doc1_lines = doc1_text.splitlines()
    doc2_lines = doc2_text.splitlines()
    
    # Use difflib to get differences
    differ = difflib.unified_diff(doc1_lines, doc2_lines, lineterm='')
    differences = []
    
    line_num = 0
    for line in differ:
        if line.startswith('---') or line.startswith('+++') or line.startswith('@@'):
            continue
        
        if line.startswith('-'):
            differences.append({
                'type': 'removed',
                'line_number': line_num,
                'content': line[1:],
                'context': 'template'
            })
        elif line.startswith('+'):
            differences.append({
                'type': 'added',
                'line_number': line_num,
                'content': line[1:],
                'context': 'contract'
            })
        
        line_num += 1
    
    logger.info(f"Found {len(differences)} differences between documents")
    return differences

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    logger.info("🚀 Starting Contract Analysis Backend...")
    
    # Create necessary directories
    os.makedirs("data", exist_ok=True)
    os.makedirs("data/uploads", exist_ok=True)
    os.makedirs("data/pdfs", exist_ok=True)
    os.makedirs("data/templates", exist_ok=True)
    os.makedirs("data/templates/purchase_agreement", exist_ok=True)
    os.makedirs("data/templates/rider", exist_ok=True)
    os.makedirs("data/templates/legal_comments", exist_ok=True)
    logger.info("📁 Created necessary directories")
    
    # Initialize database
    init_db()
    logger.info("🗄️ Database initialized")
    
    # Auto-seed templates if database is empty
    try:
        from auto_seed_on_startup import auto_seed_templates
        auto_seed_templates()
    except Exception as e:
        logger.error(f"❌ Auto-seeding failed: {e}")
    
    # Test inference service connection
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get("http://inference-service:9200/health", timeout=5.0)
            if response.status_code == 200:
                logger.info("✅ Inference service is healthy")
            else:
                logger.warning(f"⚠️ Inference service returned status {response.status_code}")
    except Exception as e:
        logger.warning(f"⚠️ Could not connect to inference service: {str(e)}")
    
    logger.info("✅ Backend startup complete")
    
    yield
    
    # Shutdown
    logger.info("🛑 Shutting down Contract Analysis Backend...")
    logger.info("✅ Backend shutdown complete")

app = FastAPI(
    title="LLM Fine-tuning Platform API",
    description="Backend API for fine-tuning and inference with LLM models",
    version="1.0.0",
    lifespan=lifespan
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://frontend:3000", "http://localhost:9000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Service URLs
PDF_PROCESSOR_URL = os.getenv("PDF_PROCESSOR_URL", "http://localhost:8001")
TRAINER_URL = os.getenv("TRAINER_URL", "http://localhost:8002")
INFERENCE_URL = os.getenv("INFERENCE_SERVICE_URL", "http://host.docker.internal:9200")

@app.get("/")
async def root():
    return {"message": "LLM Fine-tuning Platform API", "version": "1.0.0"}

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    services_status = {}
    
    async with httpx.AsyncClient() as client:
        # Check PDF processor
        try:
            response = await client.get(f"{PDF_PROCESSOR_URL}/health", timeout=5.0)
            services_status["pdf_processor"] = "healthy" if response.status_code == 200 else "unhealthy"
        except:
            services_status["pdf_processor"] = "unreachable"
        
        # Check trainer
        try:
            response = await client.get(f"{TRAINER_URL}/health", timeout=5.0)
            services_status["trainer"] = "healthy" if response.status_code == 200 else "unhealthy"
        except:
            services_status["trainer"] = "unreachable"
        
        # Check inference service
        try:
            response = await client.get(f"{INFERENCE_URL}/health", timeout=5.0)
            services_status["inference"] = "healthy" if response.status_code == 200 else "unhealthy"
        except:
            services_status["inference"] = "unreachable"
    
    return {"status": "healthy", "services": services_status}

# Document Management Endpoints
@app.post("/api/pdfs/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a PDF or DOCX file for processing"""
    if not is_supported_file(file.filename):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are allowed")
    
    file_id = str(uuid.uuid4())
    file_path = f"data/pdfs/{file_id}_{file.filename}"
    
    # Save file
    async with aiofiles.open(file_path, 'wb') as f:
        content = await file.read()
        await f.write(content)
    
    # Save to database
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO pdf_files (id, filename, file_path, uploaded_at, size_bytes)
        VALUES (?, ?, ?, ?, ?)
    ''', (file_id, file.filename, file_path, datetime.now().isoformat(), len(content)))
    conn.commit()
    conn.close()
    
    return {"id": file_id, "filename": file.filename, "status": "uploaded"}

@app.get("/api/pdfs")
async def list_pdfs():
    """List all uploaded PDF files"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM pdf_files ORDER BY uploaded_at DESC')
    rows = cursor.fetchall()
    conn.close()
    
    pdfs = []
    for row in rows:
        pdfs.append({
            "id": row[0],
            "filename": row[1],
            "file_path": row[2],
            "processed": bool(row[3]),
            "uploaded_at": row[4],
            "size_bytes": row[5]
        })
    
    return {"pdfs": pdfs}

@app.post("/api/pdfs/{pdf_id}/process")
async def process_pdf(pdf_id: str, background_tasks: BackgroundTasks):
    """Process a PDF file to extract text for training"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT file_path FROM pdf_files WHERE id = ?', (pdf_id,))
    result = cursor.fetchone()
    conn.close()
    
    if not result:
        raise HTTPException(status_code=404, detail="PDF not found")
    
    file_path = result[0]
    
    # Call PDF processor service
    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                f"{PDF_PROCESSOR_URL}/process",
                json={"file_path": file_path, "pdf_id": pdf_id}
            )
            if response.status_code == 200:
                # Update database
                conn = sqlite3.connect(DATABASE_PATH)
                cursor = conn.cursor()
                cursor.execute('UPDATE pdf_files SET processed = TRUE WHERE id = ?', (pdf_id,))
                conn.commit()
                conn.close()
                
                return {"status": "processing", "pdf_id": pdf_id}
            else:
                raise HTTPException(status_code=500, detail="PDF processing failed")
        except httpx.RequestError:
            raise HTTPException(status_code=503, detail="PDF processor service unavailable")

# Model Management Endpoints
@app.get("/api/models/available")
async def get_available_models():
    """Get list of available models for fine-tuning"""
    models = [
        {"id": "llama-3.2-1b", "name": "Llama 3.2 1B", "type": "llama", "size": "1B parameters"},
        {"id": "llama-3.2-3b", "name": "Llama 3.2 3B", "type": "llama", "size": "3B parameters"},
        {"id": "deepseek-coder-1.3b", "name": "DeepSeek Coder 1.3B", "type": "deepseek", "size": "1.3B parameters"},
        {"id": "mistral-7b", "name": "Mistral 7B", "type": "mistral", "size": "7B parameters"},
        {"id": "phi-3-mini", "name": "Phi-3 Mini", "type": "phi", "size": "3.8B parameters"},
    ]
    return {"models": models}

@app.get("/api/models/trained")
async def get_trained_models():
    """Get list of trained/fine-tuned models"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM models ORDER BY created_at DESC')
    rows = cursor.fetchall()
    conn.close()
    
    models = []
    for row in rows:
        models.append({
            "id": row[0],
            "name": row[1],
            "type": row[2],
            "status": row[3],
            "size_mb": row[4],
            "created_at": row[5],
            "file_path": row[6]
        })
    
    return {"models": models}

# Training Endpoints
@app.post("/api/training/start")
async def start_training(
    model_name: str,
    dataset_files: List[str],
    hyperparameters: dict = None
):
    """Start a fine-tuning job"""
    if hyperparameters is None:
        hyperparameters = {
            "learning_rate": 2e-5,
            "batch_size": 4,
            "num_epochs": 3,
            "max_length": 512
        }
    
    job_id = str(uuid.uuid4())
    
    # Save job to database
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO training_jobs (id, model_name, dataset_files, status, created_at, hyperparameters)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (
        job_id,
        model_name,
        json.dumps(dataset_files),
        "queued",
        datetime.now().isoformat(),
        json.dumps(hyperparameters)
    ))
    conn.commit()
    conn.close()
    
    # Call trainer service
    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                f"{TRAINER_URL}/train",
                json={
                    "job_id": job_id,
                    "model_name": model_name,
                    "dataset_files": dataset_files,
                    "hyperparameters": hyperparameters
                }
            )
            if response.status_code == 200:
                return {"job_id": job_id, "status": "started"}
            else:
                raise HTTPException(status_code=500, detail="Training service error")
        except httpx.RequestError:
            raise HTTPException(status_code=503, detail="Training service unavailable")

@app.get("/api/training/jobs")
async def get_training_jobs():
    """Get all training jobs"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM training_jobs ORDER BY created_at DESC')
    rows = cursor.fetchall()
    conn.close()
    
    jobs = []
    for row in rows:
        jobs.append({
            "id": row[0],
            "model_name": row[1],
            "dataset_files": json.loads(row[2]),
            "status": row[3],
            "created_at": row[4],
            "completed_at": row[5],
            "hyperparameters": json.loads(row[6]),
            "progress": row[7]
        })
    
    return {"jobs": jobs}

@app.get("/api/training/jobs/{job_id}")
async def get_training_job(job_id: str):
    """Get specific training job details"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM training_jobs WHERE id = ?', (job_id,))
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        raise HTTPException(status_code=404, detail="Training job not found")
    
    return {
        "id": row[0],
        "model_name": row[1],
        "dataset_files": json.loads(row[2]),
        "status": row[3],
        "created_at": row[4],
        "completed_at": row[5],
        "hyperparameters": json.loads(row[6]),
        "progress": row[7]
    }

# Inference Endpoints
@app.post("/api/inference/generate")
async def generate_text(request: InferenceRequest):
    """Generate text using a trained model"""
    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                f"{INFERENCE_URL}/generate",
                json=request.dict()
            )
            if response.status_code == 200:
                return response.json()
            else:
                raise HTTPException(status_code=500, detail="Inference service error")
        except httpx.RequestError:
            raise HTTPException(status_code=503, detail="Inference service unavailable")

@app.get("/api/inference/models")
async def get_inference_models():
    """Get available models for inference"""
    async with httpx.AsyncClient() as client:
        try:
            response = await client.get(f"{INFERENCE_URL}/models")
            if response.status_code == 200:
                return response.json()
            else:
                raise HTTPException(status_code=500, detail="Inference service error")
        except httpx.RequestError:
            raise HTTPException(status_code=503, detail="Inference service unavailable")

# Contract Review Endpoints
@app.post("/api/contracts/upload")
async def upload_contract(file: UploadFile = File(...)):
    """Upload a contract PDF or DOCX file for review"""
    logger.info(f"Starting contract upload for file: {file.filename}")
    
    if not is_supported_file(file.filename):
        logger.error(f"Invalid file type uploaded: {file.filename}")
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are allowed")
    
    file_id = str(uuid.uuid4())
    file_path = f"data/pdfs/{file_id}_{file.filename}"
    
    logger.info(f"Saving file to: {file_path}")
    
    # Save file
    async with aiofiles.open(file_path, 'wb') as f:
        content = await file.read()
        await f.write(content)
    
    logger.info(f"File saved successfully, size: {len(content)} bytes")
    
    # Extract text from file (PDF or DOCX)
    try:
        logger.info(f"Starting text extraction for uploaded file: {file.filename}")
        extracted_text = extract_text_from_file(file_path)
        text_length = len(extracted_text)
        logger.info(f"Text extraction successful. Extracted {text_length} characters")
        
        # Detect law firm from extracted text
        law_firm_id, law_firm_name = detect_law_firm(extracted_text) if extracted_text else (None, None)
        if law_firm_name:
            logger.info(f"Law firm detected: {law_firm_name}")
        else:
            logger.info("No law firm detected in contract")
        
        # Log text content summary
        if extracted_text:
            lines = extracted_text.split('\n')
            non_empty_lines = [line.strip() for line in lines if line.strip()]
            logger.info(f"Text extraction summary:")
            logger.info(f"  - Total lines: {len(lines)}")
            logger.info(f"  - Non-empty lines: {len(non_empty_lines)}")
            logger.info(f"  - Average line length: {sum(len(line) for line in non_empty_lines) / len(non_empty_lines) if non_empty_lines else 0:.1f}")
            
            # Log first few lines for debugging
            preview_lines = non_empty_lines[:5]
            logger.info(f"First few lines of extracted text:")
            for i, line in enumerate(preview_lines):
                logger.info(f"  Line {i+1}: {line[:100]}{'...' if len(line) > 100 else ''}")
        
    except Exception as e:
        logger.error(f"Text extraction failed for {file.filename}: {str(e)}")
        # If text extraction fails, still save the file but mark as not processed
        extracted_text = None
        text_length = 0
        law_firm_id, law_firm_name = None, None

    # Save to database
    logger.info(f"Saving file metadata to database")
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO pdf_files (id, filename, file_path, uploaded_at, size_bytes, extracted_text, processed, law_firm_id, law_firm_name)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (file_id, file.filename, file_path, datetime.now().isoformat(), len(content), extracted_text, extracted_text is not None, law_firm_id, law_firm_name))
    conn.commit()
    conn.close()
    
    logger.info(f"Contract upload completed successfully for {file.filename}")

    return {
        "id": file_id, 
        "filename": file.filename, 
        "status": "uploaded",
        "text_extracted": extracted_text is not None,
        "text_length": text_length,
        "law_firm_id": law_firm_id,
        "law_firm_name": law_firm_name or "N.A."
    }

@app.get("/api/contracts/files")
async def get_contract_files():
    """Get all uploaded contract files"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT id, filename, file_path, processed, uploaded_at, 
               size_bytes, extracted_text, document_type, law_firm_id, law_firm_name
        FROM pdf_files 
        ORDER BY uploaded_at DESC
    ''')
    files = cursor.fetchall()
    conn.close()
    
    return {
        "files": [
            {
                "id": file[0],
                "filename": file[1],
                "file_path": file[2],
                "processed": bool(file[3]),
                "uploaded_at": file[4],
                "size_bytes": file[5],
                "text_extracted": bool(file[6]),
                "text_length": len(file[6]) if file[6] else 0,
                "document_type": file[7],
                "law_firm_id": file[8],
                "law_firm_name": file[9]
            }
            for file in files
        ]
    }

@app.put("/api/contracts/files/{file_id}/document-type")
async def update_document_type(file_id: str, document_type: str = Body(..., embed=True)):
    """Update the document type for a contract file"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    # Check if file exists
    cursor.execute('SELECT id FROM pdf_files WHERE id = ?', (file_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Contract file not found")
    
    # Update document type
    cursor.execute('UPDATE pdf_files SET document_type = ? WHERE id = ?', (document_type, file_id))
    conn.commit()
    conn.close()
    
    return {"message": "Document type updated successfully", "file_id": file_id, "document_type": document_type}

@app.post("/api/contracts/review")
async def review_contracts(request: ContractReviewRequest):
    """Review contract files using the fine-tuned model"""
    review_id = str(uuid.uuid4())
    
    logger.info(f"Starting contract review {review_id}")
    logger.info(f"Request details: files={request.contract_files}, model={request.model_id}, type={request.review_type}")
    
    # Save review request to database
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO contract_reviews (id, contract_files, model_id, review_type, status, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (review_id, json.dumps(request.contract_files), request.model_id, request.review_type, "processing", datetime.now().isoformat()))
    conn.commit()
    
    try:
        # Get contract texts
        contract_sections = []
        logger.info(f"Retrieving contract texts for {len(request.contract_files)} files")
        
        for i, file_id in enumerate(request.contract_files):
            logger.info(f"Processing contract file {i+1}/{len(request.contract_files)}: {file_id}")
            cursor.execute('SELECT extracted_text, filename FROM pdf_files WHERE id = ?', (file_id,))
            row = cursor.fetchone()
            
            if row and row[0]:
                filename = row[1]
                text_content = row[0]
                contract_text = f"Contract: {filename}\n\n{text_content}"
                contract_sections.append(contract_text)
                
                logger.info(f"Retrieved contract text for {filename}")
                logger.info(f"  - Text length: {len(text_content)} characters")
                logger.info(f"  - Preview: {text_content[:200]}...")
                
            else:
                logger.error(f"Contract file {file_id} not found or text not extracted")
                raise HTTPException(status_code=400, detail=f"Contract file {file_id} not found or text not extracted")
        
        conn.close()
        
        total_text_length = sum(len(section) for section in contract_sections)
        logger.info(f"Prepared {len(contract_sections)} contract sections for review")
        logger.info(f"Total text length: {total_text_length} characters")
        
        # Call inference service for contract review
        logger.info(f"Calling inference service at {INFERENCE_URL}/contract/review")
        async with httpx.AsyncClient(timeout=300.0) as client:  # 5 minute timeout
            start_time = datetime.now()
            
            inference_request = {
                "contract_sections": contract_sections,
                "model_id": request.model_id,
                "review_type": request.review_type
            }
            
            logger.info(f"Sending inference request:")
            logger.info(f"  - Model ID: {request.model_id}")
            logger.info(f"  - Review type: {request.review_type}")
            logger.info(f"  - Number of sections: {len(contract_sections)}")
            
            response = await client.post(
                f"{INFERENCE_URL}/contract/review",
                json=inference_request
            )
            
            end_time = datetime.now()
            processing_time = (end_time - start_time).total_seconds()
            
            logger.info(f"Inference service response received in {processing_time:.2f} seconds")
            logger.info(f"Response status: {response.status_code}")
            
            if response.status_code == 200:
                review_result = response.json()
                logger.info(f"Contract review completed successfully")
                logger.info(f"Review result keys: {list(review_result.keys())}")
                
                # Update database with results
                conn = sqlite3.connect(DATABASE_PATH)
                cursor = conn.cursor()
                cursor.execute('''
                    UPDATE contract_reviews 
                    SET status = ?, completed_at = ?, review_result = ?, processing_time = ?
                    WHERE id = ?
                ''', ("completed", end_time.isoformat(), json.dumps(review_result), processing_time, review_id))
                conn.commit()
                conn.close()
                
                logger.info(f"Contract review {review_id} completed and saved to database")
                
                return {
                    "review_id": review_id,
                    "status": "completed",
                    "review": review_result,
                    "processing_time": processing_time
                }
            else:
                logger.error(f"Inference service returned error: {response.status_code}")
                logger.error(f"Response content: {response.text}")
                
                # Update database with error
                conn = sqlite3.connect(DATABASE_PATH)
                cursor = conn.cursor()
                cursor.execute('''
                    UPDATE contract_reviews 
                    SET status = ?, completed_at = ?
                    WHERE id = ?
                ''', ("failed", datetime.now().isoformat(), review_id))
                conn.commit()
                conn.close()
                
                raise HTTPException(status_code=500, detail="Contract review failed")
                
    except httpx.RequestError as e:
        logger.error(f"Request error when calling inference service: {str(e)}")
        # Update database with error
        conn = sqlite3.connect(DATABASE_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            UPDATE contract_reviews 
            SET status = ?, completed_at = ?
            WHERE id = ?
        ''', ("failed", datetime.now().isoformat(), review_id))
        conn.commit()
        conn.close()
        
        raise HTTPException(status_code=503, detail="Inference service unavailable")

@app.get("/api/contracts/reviews")
async def get_contract_reviews():
    """Get all contract review history"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM contract_reviews ORDER BY created_at DESC')
    rows = cursor.fetchall()
    conn.close()
    
    reviews = []
    for row in rows:
        review_result = json.loads(row[7]) if row[7] else None
        reviews.append({
            "id": row[0],
            "contract_files": json.loads(row[1]),
            "model_id": row[2],
            "review_type": row[3],
            "status": row[4],
            "created_at": row[5],
            "completed_at": row[6],
            "review_result": review_result,
            "processing_time": row[8]
        })
    
    return {"reviews": reviews}

@app.get("/api/contracts/reviews/{review_id}")
async def get_contract_review(review_id: str):
    """Get specific contract review details"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM contract_reviews WHERE id = ?', (review_id,))
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        raise HTTPException(status_code=404, detail="Contract review not found")
    
    review_result = json.loads(row[7]) if row[7] else None
    return {
        "id": row[0],
        "contract_files": json.loads(row[1]),
        "model_id": row[2],
        "review_type": row[3],
        "status": row[4],
        "created_at": row[5],
        "completed_at": row[6],
        "review_result": review_result,
        "processing_time": row[8]
    }

# Law Firm Management Endpoints
@app.get("/api/law-firms")
async def get_law_firms():
    """Get all law firms"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT id, name, keywords, created_at FROM law_firms ORDER BY name')
    rows = cursor.fetchall()
    conn.close()
    
    law_firms = []
    for row in rows:
        law_firms.append({
            "id": row[0],
            "name": row[1],
            "keywords": json.loads(row[2]),
            "created_at": row[3]
        })
    
    return {"law_firms": law_firms}

@app.post("/api/law-firms")
async def create_law_firm(law_firm: LawFirmCreate):
    """Create a new law firm"""
    law_firm_id = str(uuid.uuid4())
    
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute('''
            INSERT INTO law_firms (id, name, keywords, created_at)
            VALUES (?, ?, ?, ?)
        ''', (law_firm_id, law_firm.name, json.dumps(law_firm.keywords), datetime.now().isoformat()))
        conn.commit()
    except sqlite3.IntegrityError:
        conn.close()
        raise HTTPException(status_code=400, detail="Law firm with this name already exists")
    finally:
        conn.close()
    
    return {"id": law_firm_id, "name": law_firm.name, "keywords": law_firm.keywords, "status": "created"}

# Template Management Endpoints
@app.get("/api/templates")
async def get_templates(law_firm_id: Optional[str] = None, template_type: Optional[str] = None):
    """Get contract templates"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    query = 'SELECT id, law_firm_id, law_firm_name, template_type, file_name, file_path, created_at FROM contract_templates'
    params = []
    conditions = []
    
    if law_firm_id:
        conditions.append('law_firm_id = ?')
        params.append(law_firm_id)
    
    if template_type:
        conditions.append('template_type = ?')
        params.append(template_type)
    
    if conditions:
        query += ' WHERE ' + ' AND '.join(conditions)
    
    query += ' ORDER BY law_firm_name, template_type'
    
    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()
    
    templates = []
    for row in rows:
        # Get file size
        file_size = 0
        try:
            file_path = Path(row[5])
            if file_path.exists():
                file_size = file_path.stat().st_size
        except Exception:
            pass
            
        templates.append({
            "id": row[0],
            "law_firm_id": row[1],
            "law_firm_name": row[2],
            "template_type": row[3],
            "filename": row[4],  # Changed from file_name to filename for frontend consistency
            "file_path": row[5],
            "file_size": file_size,
            "created_at": row[6]
        })
    
    return {"templates": templates}

@app.post("/api/templates/upload")
async def upload_template(
    law_firm_id: str,
    template_type: str,
    file: UploadFile = File(...)
):
    """Upload a contract template (PDF or DOCX)"""
    logger.info(f"Template upload request: law_firm_id={law_firm_id}, template_type={template_type}, filename={file.filename}")
    
    if not is_supported_file(file.filename):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are allowed")
    
    if template_type not in ['purchase_agreement', 'rider', 'legal_comments']:
        raise HTTPException(status_code=400, detail="Invalid template type")
    
    # Get law firm name
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT name FROM law_firms WHERE id = ?', (law_firm_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Law firm not found")
    law_firm_name = row[0]
    
    template_id = str(uuid.uuid4())
    
    # Create directory structure if it doesn't exist
    template_dir = f"data/templates/{template_type}"
    os.makedirs(template_dir, exist_ok=True)
    
    file_path = f"{template_dir}/{template_id}_{file.filename}"
    
    # Save file
    try:
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        logger.info(f"Template file saved to: {file_path}")
    except Exception as e:
        logger.error(f"Failed to save template file: {str(e)}")
        conn.close()
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
    
    # Extract text content
    extracted_text = ""
    try:
        extracted_text = extract_text_from_file(file_path)
        logger.info(f"Extracted {len(extracted_text)} characters from template")
    except Exception as e:
        logger.warning(f"Could not extract text from template: {str(e)}")
    
    # Save to database
    try:
        cursor.execute('''
            INSERT INTO contract_templates (id, law_firm_id, law_firm_name, template_type, file_name, file_path, content, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (template_id, law_firm_id, law_firm_name, template_type, file.filename, file_path, extracted_text, datetime.now().isoformat()))
        conn.commit()
        logger.info(f"Template saved to database with ID: {template_id}")
    except Exception as e:
        logger.error(f"Failed to save template to database: {str(e)}")
        conn.close()
        # Clean up the file if database save failed
        try:
            os.remove(file_path)
        except:
            pass
        raise HTTPException(status_code=500, detail=f"Failed to save to database: {str(e)}")
    finally:
        conn.close()
    
    return {
        "id": template_id,
        "law_firm_id": law_firm_id,
        "law_firm_name": law_firm_name,
        "template_type": template_type,
        "filename": file.filename,
        "status": "uploaded"
    }

@app.get("/api/templates/{template_id}/content")
async def get_template_content(template_id: str):
    """Get template content"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT content, template_type, law_firm_name FROM contract_templates WHERE id = ?', (template_id,))
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        raise HTTPException(status_code=404, detail="Template not found")
    
    return {
        "template_id": template_id,
        "content": row[0],
        "template_type": row[1],
        "law_firm_name": row[2]
    }

@app.get("/api/templates/{template_id}/download")
async def download_template(template_id: str):
    """Download template file"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT file_path, file_name FROM contract_templates WHERE id = ?', (template_id,))
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        raise HTTPException(status_code=404, detail="Template not found")
    
    file_path, file_name = row
    full_path = Path(file_path)
    
    if not full_path.exists():
        raise HTTPException(status_code=404, detail="Template file not found on disk")
    
    return FileResponse(
        path=str(full_path),
        filename=file_name,
        media_type='application/octet-stream'
    )

@app.delete("/api/templates/{template_id}")
async def delete_template(template_id: str):
    """Delete a template"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    # Get template info before deletion
    cursor.execute('SELECT file_path, file_name FROM contract_templates WHERE id = ?', (template_id,))
    row = cursor.fetchone()
    
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Template not found")
    
    file_path, file_name = row
    
    # Delete from database
    cursor.execute('DELETE FROM contract_templates WHERE id = ?', (template_id,))
    deleted_rows = cursor.rowcount
    conn.commit()
    conn.close()
    
    if deleted_rows > 0:
        # Try to delete file from disk
        try:
            full_path = Path(file_path)
            if full_path.exists():
                full_path.unlink()
                logger.info(f"Deleted template file: {file_path}")
        except Exception as e:
            logger.warning(f"Could not delete template file {file_path}: {e}")
        
        return {"message": f"Template '{file_name}' deleted successfully"}
    else:
        raise HTTPException(status_code=404, detail="Template not found")

# Contract Comment Endpoints
@app.post("/api/test-endpoint")
async def test_endpoint():
    """Test endpoint to verify routing works"""
    return {"test": "POST endpoint working", "timestamp": datetime.now().isoformat()}

@app.post("/api/contracts/comment/batch")
async def generate_batch_contract_comments(request: BatchContractCommentRequest):
    """Generate AI comments for multiple contracts using law firm templates - one overall comment with individual diffs"""
    comment_id = str(uuid.uuid4())
    start_time = datetime.now()
    
    logger.info(f"Starting batch contract comment generation for {len(request.contract_file_ids)} contracts")
    
    # Save comment request to database
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    # Use the first contract file ID as the primary reference for the comment record
    primary_contract_id = request.contract_file_ids[0]
    
    cursor.execute('''
        INSERT INTO contract_comments (id, contract_file_id, law_firm_id, status, created_at)
        VALUES (?, ?, ?, ?, ?)
    ''', (comment_id, primary_contract_id, request.law_firm_id, "processing", start_time.isoformat()))
    conn.commit()
    
    try:
        # Get law firm info
        cursor.execute('SELECT name FROM law_firms WHERE id = ?', (request.law_firm_id,))
        law_firm_row = cursor.fetchone()
        if not law_firm_row:
            conn.close()
            raise HTTPException(status_code=404, detail="Law firm not found")
        
        law_firm_name = law_firm_row[0]
        
        # Get templates for this law firm
        cursor.execute('''
            SELECT id, template_type, file_name, content 
            FROM contract_templates 
            WHERE law_firm_id = ?
        ''', (request.law_firm_id,))
        template_rows = cursor.fetchall()
        
        if not template_rows:
            conn.close()
            raise HTTPException(status_code=404, detail=f"No templates found for law firm: {law_firm_name}")
        
        # Parse legal comments from templates (only once for all documents)
        comments = []
        legal_comments_content = None
        legal_comments_template_id = None
        legal_comments_template_filename = None
        
        for template_row in template_rows:
            template_id, template_type, template_filename, template_content = template_row
            if template_type == 'legal_comments':
                legal_comments_content = template_content
                legal_comments_template_id = template_id
                legal_comments_template_filename = template_filename
                break
        
        if legal_comments_content:
            logger.info(f"Parsing legal comments content, length: {len(legal_comments_content)}")
            comments = parse_legal_comments(legal_comments_content)
            logger.info(f"Parsed {len(comments)} comments")
        else:
            logger.warning("No legal_comments template found or template content is empty")
        
        # Process each contract file individually for comparisons
        all_comparisons = []
        contract_files_info = []
        
        for contract_file_id in request.contract_file_ids:
            # Get contract file info
            cursor.execute('SELECT filename, extracted_text, document_type FROM pdf_files WHERE id = ?', (contract_file_id,))
            contract_row = cursor.fetchone()
            if not contract_row:
                logger.warning(f"Contract file not found: {contract_file_id}")
                continue
            
            contract_filename, contract_text, contract_document_type = contract_row
            contract_files_info.append({
                "file_id": contract_file_id,
                "filename": contract_filename,
                "document_type": contract_document_type
            })
            
            # Generate template comparisons based on document type
            contract_comparisons = []
            
            # If contract has a document type, only compare with matching template type
            if contract_document_type:
                logger.info(f"Contract {contract_filename} document type: {contract_document_type}")
                matching_templates = [t for t in template_rows if t[1] == contract_document_type]
                if not matching_templates:
                    logger.warning(f"No templates found for document type: {contract_document_type}")
            else:
                # If no document type specified, compare with all templates except legal_comments
                matching_templates = [t for t in template_rows if t[1] != 'legal_comments']
            
            for template_row in matching_templates:
                template_id, template_type, template_filename, template_content = template_row
                
                logger.info(f"Starting document comparison for {contract_filename}...")
                differences = compare_documents(template_content, contract_text)
                logger.info(f"Found {len(differences)} differences between {contract_filename} and {template_filename}")
                
                contract_comparisons.append({
                    "comparison_id": str(uuid.uuid4()),
                    "template_id": template_id,
                    "template_type": template_type,
                    "template_filename": template_filename,
                    "differences": differences
                })
            
            all_comparisons.append({
                "contract_file_id": contract_file_id,
                "contract_filename": contract_filename,
                "document_type": contract_document_type,
                "comparisons": contract_comparisons
            })
        
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        result = {
            "comment_id": comment_id,
            "contract_files": contract_files_info,
            "law_firm_id": request.law_firm_id,
            "law_firm_name": law_firm_name,
            "comments": comments,  # Single set of legal comments for all documents
            "legal_comments_template_id": legal_comments_template_id,  # Template ID for download
            "legal_comments_template_filename": legal_comments_template_filename,  # Template filename
            "document_comparisons": all_comparisons,  # Individual comparisons for each document
            "template_comparison": None  # Legacy field for backward compatibility
        }
        
        # Update database with results
        cursor.execute('''
            UPDATE contract_comments 
            SET status = ?, completed_at = ?, comments_result = ?, processing_time = ?
            WHERE id = ?
        ''', ("completed", end_time.isoformat(), json.dumps(result), processing_time, comment_id))
        conn.commit()
        conn.close()
        
        logger.info(f"Batch contract comment generation completed for {len(request.contract_file_ids)} files")
        
        return result
        
    except Exception as e:
        logger.error(f"Error in batch contract comment generation: {str(e)}")
        
        # Create a new database connection for error handling
        try:
            error_conn = sqlite3.connect(DATABASE_PATH)
            error_cursor = error_conn.cursor()
            error_cursor.execute('''
                UPDATE contract_comments 
                SET status = ?, completed_at = ?
                WHERE id = ?
            ''', ("failed", datetime.now().isoformat(), comment_id))
            error_conn.commit()
            error_conn.close()
        except Exception as db_error:
            logger.error(f"Failed to update database with error status: {str(db_error)}")
        
        # Ensure the original connection is closed if it's still open
        try:
            if 'conn' in locals() and conn:
                conn.close()
        except:
            pass
        
        raise HTTPException(status_code=500, detail=f"Batch comment generation failed: {str(e)}")

@app.post("/api/contracts/comment")
async def generate_contract_comments(request: ContractCommentRequest):
    """Generate AI comments for a contract using law firm templates"""
    comment_id = str(uuid.uuid4())
    start_time = datetime.now()
    
    logger.info(f"Starting contract comment generation for contract: {request.contract_file_id}")
    
    # Save comment request to database
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO contract_comments (id, contract_file_id, law_firm_id, status, created_at)
        VALUES (?, ?, ?, ?, ?)
    ''', (comment_id, request.contract_file_id, request.law_firm_id, "processing", start_time.isoformat()))
    conn.commit()
    
    try:
        # Get contract file info
        cursor.execute('SELECT filename, extracted_text, document_type FROM pdf_files WHERE id = ?', (request.contract_file_id,))
        contract_row = cursor.fetchone()
        if not contract_row:
            conn.close()
            raise HTTPException(status_code=404, detail="Contract file not found")
        
        contract_filename, contract_text, contract_document_type = contract_row
        
        # Get law firm info
        cursor.execute('SELECT name FROM law_firms WHERE id = ?', (request.law_firm_id,))
        law_firm_row = cursor.fetchone()
        if not law_firm_row:
            conn.close()
            raise HTTPException(status_code=404, detail="Law firm not found")
        
        law_firm_name = law_firm_row[0]
        
        # Get templates for this law firm
        cursor.execute('''
            SELECT id, template_type, file_name, content 
            FROM contract_templates 
            WHERE law_firm_id = ?
        ''', (request.law_firm_id,))
        template_rows = cursor.fetchall()
        
        if not template_rows:
            conn.close()
            raise HTTPException(status_code=404, detail=f"No templates found for law firm: {law_firm_name}")
        
        # Parse legal comments from templates (only once for all documents)
        comments = []
        legal_comments_content = None
        legal_comments_template_id = None
        legal_comments_template_filename = None
        
        for template_row in template_rows:
            template_id, template_type, template_filename, template_content = template_row
            if template_type == 'legal_comments':
                legal_comments_content = template_content
                legal_comments_template_id = template_id
                legal_comments_template_filename = template_filename
                break
        
        if legal_comments_content:
            logger.info(f"Parsing legal comments content, length: {len(legal_comments_content)}")
            comments = parse_legal_comments(legal_comments_content)
            logger.info(f"Parsed {len(comments)} comments")
        else:
            logger.warning("No legal_comments template found or template content is empty")
        
        # Generate template comparisons based on document type
        comparisons = []
        
        # If contract has a document type, only compare with matching template type
        if contract_document_type:
            logger.info(f"Contract document type: {contract_document_type}")
            matching_templates = [t for t in template_rows if t[1] == contract_document_type]
            if not matching_templates:
                logger.warning(f"No templates found for document type: {contract_document_type}")
        else:
            # If no document type specified, compare with all templates except legal_comments
            matching_templates = [t for t in template_rows if t[1] != 'legal_comments']
        
        for template_row in matching_templates:
            template_id, template_type, template_filename, template_content = template_row
            
            logger.info(f"Starting document comparison...")
            differences = compare_documents(template_content, contract_text)
            logger.info(f"Found {len(differences)} differences between documents")
            
            comparisons.append({
                "comparison_id": str(uuid.uuid4()),
                "template_id": template_id,
                "template_type": template_type,
                "template_filename": template_filename,
                "differences": differences
            })
        
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        result = {
            "comment_id": comment_id,
            "contract_file_id": request.contract_file_id,
            "contract_filename": contract_filename,
            "law_firm_id": request.law_firm_id,
            "law_firm_name": law_firm_name,
            "comparisons": comparisons,
            "comments": comments,
            "legal_comments_template_id": legal_comments_template_id,  # Template ID for download
            "legal_comments_template_filename": legal_comments_template_filename,  # Template filename
            "template_comparison": None  # Legacy field for backward compatibility
        }
        
        # Update database with results
        cursor.execute('''
            UPDATE contract_comments 
            SET status = ?, completed_at = ?, comments_result = ?, processing_time = ?
            WHERE id = ?
        ''', ("completed", end_time.isoformat(), json.dumps(result), processing_time, comment_id))
        conn.commit()
        conn.close()
        
        logger.info(f"Contract comment generation completed for {contract_filename}")
        
        return result
        
    except Exception as e:
        logger.error(f"Error in contract comment generation: {str(e)}")
        
        # Create a new database connection for error handling
        try:
            error_conn = sqlite3.connect(DATABASE_PATH)
            error_cursor = error_conn.cursor()
            error_cursor.execute('''
                UPDATE contract_comments 
                SET status = ?, completed_at = ?
                WHERE id = ?
            ''', ("failed", datetime.now().isoformat(), comment_id))
            error_conn.commit()
            error_conn.close()
        except Exception as db_error:
            logger.error(f"Failed to update database with error status: {str(db_error)}")
        
        # Ensure the original connection is closed if it's still open
        try:
            if 'conn' in locals() and conn:
                conn.close()
        except:
            pass
        
        raise HTTPException(status_code=500, detail=f"Comment generation failed: {str(e)}")

@app.get("/api/contracts/comments")
async def get_contract_comments():
    """Get all contract comment history"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT cc.*, pf.filename, lf.name as law_firm_name 
        FROM contract_comments cc
        LEFT JOIN pdf_files pf ON cc.contract_file_id = pf.id
        LEFT JOIN law_firms lf ON cc.law_firm_id = lf.id
        ORDER BY cc.created_at DESC
    ''')
    rows = cursor.fetchall()
    conn.close()
    
    comments = []
    for row in rows:
        comments_result = json.loads(row[6]) if row[6] else None
        comments.append({
            "id": row[0],
            "contract_file_id": row[1],
            "law_firm_id": row[2],
            "status": row[3],
            "created_at": row[4],
            "completed_at": row[5],
            "comments_result": comments_result,
            "processing_time": row[7],
            "contract_filename": row[8],
            "law_firm_name": row[9]
        })
    
    return {"comments": comments}

@app.get("/api/contracts/comments/{comment_id}")
async def get_contract_comment(comment_id: str):
    """Get specific contract comment details"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT cc.*, pf.filename, lf.name as law_firm_name 
        FROM contract_comments cc
        LEFT JOIN pdf_files pf ON cc.contract_file_id = pf.id
        LEFT JOIN law_firms lf ON cc.law_firm_id = lf.id
        WHERE cc.id = ?
    ''', (comment_id,))
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        raise HTTPException(status_code=404, detail="Contract comment not found")
    
    comments_result = json.loads(row[6]) if row[6] else None
    return {
        "id": row[0],
        "contract_file_id": row[1],
        "law_firm_id": row[2],
        "status": row[3],
        "created_at": row[4],
        "completed_at": row[5],
        "comments_result": comments_result,
        "processing_time": row[7],
        "contract_filename": row[8],
        "law_firm_name": row[9]
    }

def parse_legal_comments(text: str) -> List[Dict]:
    """Parse legal comments from template content - return original content without artificial annotations"""
    logger.info(f"Starting to parse legal comments, text length: {len(text)}")
    
    # Simply return the original legal comments content as-is
    # No artificial severity levels, section parsing, or annotations
    comments = [{
        "comment": text.strip(),
        "section": "Legal Comments",
        "severity": "info",  # Neutral severity
        "suggestion": None
    }]
    
    logger.info(f"Returning original legal comments content as single comment")
    return comments

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=9100) 